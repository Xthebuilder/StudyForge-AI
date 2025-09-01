#!/usr/bin/env python3
"""
StudyForge AI - RAG (Retrieval Augmented Generation) Processor
Advanced document processing and retrieval system for enhanced AI responses
"""

import asyncio
import json
import logging
import hashlib
import time
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass, asdict
import sqlite3
from datetime import datetime

# Document processing imports
try:
    import PyPDF2
    import docx
    from bs4 import BeautifulSoup
    import markdown
except ImportError:
    PyPDF2 = None
    docx = None
    BeautifulSoup = None
    markdown = None

# Text processing
import re
from collections import Counter

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of processed document content"""
    chunk_id: str
    document_id: str
    content: str
    metadata: Dict[str, Any]
    embedding_summary: Optional[str] = None
    relevance_score: float = 0.0
    created_at: str = None
    
    def __post_init__(self):
        if not self.created_at:
            self.created_at = datetime.now().isoformat()

@dataclass
class ProcessedDocument:
    """Represents a processed document with metadata"""
    document_id: str
    filename: str
    file_type: str
    file_size: int
    content_preview: str
    chunks_count: int
    keywords: List[str]
    processing_time: float
    upload_time: str
    error_message: Optional[str] = None

class DocumentProcessor:
    """Advanced document processing with multiple format support"""
    
    def __init__(self):
        self.supported_types = {
            'text/plain': self._process_text,
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'text/html': self._process_html,
            'text/markdown': self._process_markdown,
            'application/json': self._process_json,
            'text/csv': self._process_csv,
            'text/xml': self._process_xml
        }
        
    async def process_document(self, file_content: bytes, filename: str, content_type: str) -> ProcessedDocument:
        """Process a document and extract meaningful content"""
        start_time = time.time()
        document_id = self._generate_document_id(filename, file_content)
        
        try:
            # Process based on content type
            processor = self.supported_types.get(content_type, self._process_text)
            content = await processor(file_content, filename)
            
            # Create chunks
            chunks = self._create_chunks(content, document_id)
            
            # Extract keywords
            keywords = self._extract_keywords(content)
            
            # Create preview
            preview = content[:500] + "..." if len(content) > 500 else content
            
            processing_time = time.time() - start_time
            
            return ProcessedDocument(
                document_id=document_id,
                filename=filename,
                file_type=content_type,
                file_size=len(file_content),
                content_preview=preview,
                chunks_count=len(chunks),
                keywords=keywords,
                processing_time=processing_time,
                upload_time=datetime.now().isoformat()
            )
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            return ProcessedDocument(
                document_id=document_id,
                filename=filename,
                file_type=content_type,
                file_size=len(file_content),
                content_preview="",
                chunks_count=0,
                keywords=[],
                processing_time=time.time() - start_time,
                upload_time=datetime.now().isoformat(),
                error_message=str(e)
            )
    
    def _generate_document_id(self, filename: str, content: bytes) -> str:
        """Generate unique document ID"""
        content_hash = hashlib.md5(content).hexdigest()
        return f"doc_{filename}_{content_hash[:8]}"
    
    async def _process_text(self, content: bytes, filename: str) -> str:
        """Process plain text files"""
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return content.decode('latin-1')
            except UnicodeDecodeError:
                return content.decode('utf-8', errors='ignore')
    
    async def _process_pdf(self, content: bytes, filename: str) -> str:
        """Process PDF files"""
        if not PyPDF2:
            raise ImportError("PyPDF2 not available for PDF processing")
        
        import io
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1} from {filename}: {e}")
        
        return text.strip()
    
    async def _process_docx(self, content: bytes, filename: str) -> str:
        """Process DOCX files"""
        if not docx:
            raise ImportError("python-docx not available for DOCX processing")
        
        import io
        doc = docx.Document(io.BytesIO(content))
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        return "\n".join(paragraphs)
    
    async def _process_html(self, content: bytes, filename: str) -> str:
        """Process HTML files"""
        if not BeautifulSoup:
            raise ImportError("BeautifulSoup not available for HTML processing")
        
        html_content = content.decode('utf-8', errors='ignore')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text content
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        return '\n'.join(line for line in lines if line)
    
    async def _process_markdown(self, content: bytes, filename: str) -> str:
        """Process Markdown files"""
        md_content = content.decode('utf-8', errors='ignore')
        
        if markdown:
            # Convert to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
        else:
            # Basic markdown processing without dependencies
            text = md_content
            # Remove markdown syntax
            text = re.sub(r'#{1,6}\s+', '', text)  # Headers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            return text
    
    async def _process_json(self, content: bytes, filename: str) -> str:
        """Process JSON files"""
        try:
            json_content = json.loads(content.decode('utf-8'))
            return json.dumps(json_content, indent=2)
        except json.JSONDecodeError:
            return content.decode('utf-8', errors='ignore')
    
    async def _process_csv(self, content: bytes, filename: str) -> str:
        """Process CSV files"""
        csv_content = content.decode('utf-8', errors='ignore')
        lines = csv_content.split('\n')
        
        # Format as readable text
        formatted_lines = []
        for i, line in enumerate(lines[:50]):  # Limit to first 50 rows
            if line.strip():
                formatted_lines.append(f"Row {i+1}: {line}")
        
        if len(lines) > 50:
            formatted_lines.append(f"... and {len(lines) - 50} more rows")
        
        return '\n'.join(formatted_lines)
    
    async def _process_xml(self, content: bytes, filename: str) -> str:
        """Process XML files"""
        if BeautifulSoup:
            xml_content = content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(xml_content, 'xml')
            return soup.get_text()
        else:
            return content.decode('utf-8', errors='ignore')
    
    def _create_chunks(self, content: str, document_id: str, chunk_size: int = 1000, overlap: int = 200) -> List[DocumentChunk]:
        """Split content into overlapping chunks for better retrieval"""
        chunks = []
        
        # Split into sentences for better boundary handling
        sentences = re.split(r'(?<=[.!?])\s+', content)
        
        current_chunk = ""
        current_size = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            if current_size + sentence_size > chunk_size and current_chunk:
                # Save current chunk
                chunk_id = f"{document_id}_chunk_{chunk_index}"
                chunks.append(DocumentChunk(
                    chunk_id=chunk_id,
                    document_id=document_id,
                    content=current_chunk.strip(),
                    metadata={
                        'chunk_index': chunk_index,
                        'word_count': len(current_chunk.split()),
                        'char_count': len(current_chunk)
                    }
                ))
                
                # Start new chunk with overlap
                chunk_index += 1
                if overlap > 0 and len(current_chunk) > overlap:
                    current_chunk = current_chunk[-overlap:] + " " + sentence
                    current_size = len(current_chunk)
                else:
                    current_chunk = sentence
                    current_size = sentence_size
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_size += sentence_size
        
        # Add final chunk if not empty
        if current_chunk.strip():
            chunk_id = f"{document_id}_chunk_{chunk_index}"
            chunks.append(DocumentChunk(
                chunk_id=chunk_id,
                document_id=document_id,
                content=current_chunk.strip(),
                metadata={
                    'chunk_index': chunk_index,
                    'word_count': len(current_chunk.split()),
                    'char_count': len(current_chunk)
                }
            ))
        
        return chunks
    
    def _extract_keywords(self, content: str, max_keywords: int = 20) -> List[str]:
        """Extract relevant keywords from content"""
        # Clean and tokenize
        words = re.findall(r'\b[a-zA-Z]{3,}\b', content.lower())
        
        # Remove common stop words
        stop_words = {
            'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our',
            'have', 'this', 'will', 'been', 'has', 'that', 'there', 'they', 'were', 'from', 'with', 'would',
            'more', 'what', 'when', 'where', 'who', 'why', 'how', 'could', 'should', 'than', 'other', 'some'
        }
        
        filtered_words = [word for word in words if word not in stop_words and len(word) > 3]
        
        # Count frequency and return most common
        word_counts = Counter(filtered_words)
        return [word for word, count in word_counts.most_common(max_keywords)]


class RAGDatabase:
    """Database manager for RAG document storage and retrieval"""
    
    def __init__(self, db_path: str = "rag_documents.db"):
        self.db_path = Path(db_path)
        self._init_database()
    
    def _init_database(self):
        """Initialize RAG database tables"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS documents (
                    document_id TEXT PRIMARY KEY,
                    filename TEXT NOT NULL,
                    file_type TEXT NOT NULL,
                    file_size INTEGER NOT NULL,
                    content_preview TEXT,
                    chunks_count INTEGER NOT NULL,
                    keywords TEXT,  -- JSON array
                    processing_time REAL NOT NULL,
                    upload_time TEXT NOT NULL,
                    error_message TEXT
                )
            ''')
            
            conn.execute('''
                CREATE TABLE IF NOT EXISTS document_chunks (
                    chunk_id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    content TEXT NOT NULL,
                    metadata TEXT,  -- JSON
                    embedding_summary TEXT,
                    relevance_score REAL DEFAULT 0.0,
                    created_at TEXT NOT NULL,
                    FOREIGN KEY (document_id) REFERENCES documents (document_id)
                )
            ''')
            
            conn.execute('''
                CREATE INDEX IF NOT EXISTS idx_document_chunks_document_id 
                ON document_chunks(document_id)
            ''')
            
            conn.execute('''
                CREATE INDEX IF NOT EXISTS idx_document_chunks_relevance 
                ON document_chunks(relevance_score DESC)
            ''')
    
    def store_document(self, doc: ProcessedDocument, chunks: List[DocumentChunk]):
        """Store processed document and its chunks"""
        with sqlite3.connect(self.db_path) as conn:
            # Store document
            conn.execute('''
                INSERT OR REPLACE INTO documents 
                (document_id, filename, file_type, file_size, content_preview, 
                 chunks_count, keywords, processing_time, upload_time, error_message)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                doc.document_id, doc.filename, doc.file_type, doc.file_size,
                doc.content_preview, doc.chunks_count, json.dumps(doc.keywords),
                doc.processing_time, doc.upload_time, doc.error_message
            ))
            
            # Store chunks
            for chunk in chunks:
                conn.execute('''
                    INSERT OR REPLACE INTO document_chunks 
                    (chunk_id, document_id, content, metadata, embedding_summary, 
                     relevance_score, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (
                    chunk.chunk_id, chunk.document_id, chunk.content,
                    json.dumps(chunk.metadata), chunk.embedding_summary,
                    chunk.relevance_score, chunk.created_at
                ))
    
    def search_relevant_chunks(self, query: str, limit: int = 5) -> List[DocumentChunk]:
        """Search for relevant document chunks using keyword matching"""
        query_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', query.lower()))
        
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            rows = conn.execute('''
                SELECT chunk_id, document_id, content, metadata, 
                       embedding_summary, relevance_score, created_at
                FROM document_chunks
                ORDER BY relevance_score DESC
                LIMIT ?
            ''', (limit * 3,)).fetchall()  # Get more to filter
            
            chunks_with_scores = []
            for row in rows:
                chunk_content = row['content'].lower()
                
                # Calculate relevance score based on query word matches
                matches = sum(1 for word in query_words if word in chunk_content)
                score = matches / len(query_words) if query_words else 0.0
                
                if score > 0.1:  # Minimum relevance threshold
                    chunk = DocumentChunk(
                        chunk_id=row['chunk_id'],
                        document_id=row['document_id'],
                        content=row['content'],
                        metadata=json.loads(row['metadata']) if row['metadata'] else {},
                        embedding_summary=row['embedding_summary'],
                        relevance_score=score,
                        created_at=row['created_at']
                    )
                    chunks_with_scores.append(chunk)
            
            # Sort by calculated relevance and return top results
            chunks_with_scores.sort(key=lambda x: x.relevance_score, reverse=True)
            return chunks_with_scores[:limit]
    
    def get_all_documents(self) -> List[ProcessedDocument]:
        """Get all stored documents"""
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            rows = conn.execute('''
                SELECT document_id, filename, file_type, file_size, content_preview,
                       chunks_count, keywords, processing_time, upload_time, error_message
                FROM documents
                ORDER BY upload_time DESC
            ''').fetchall()
            
            documents = []
            for row in rows:
                doc = ProcessedDocument(
                    document_id=row['document_id'],
                    filename=row['filename'],
                    file_type=row['file_type'],
                    file_size=row['file_size'],
                    content_preview=row['content_preview'],
                    chunks_count=row['chunks_count'],
                    keywords=json.loads(row['keywords']) if row['keywords'] else [],
                    processing_time=row['processing_time'],
                    upload_time=row['upload_time'],
                    error_message=row['error_message']
                )
                documents.append(doc)
            
            return documents
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document and all its chunks"""
        with sqlite3.connect(self.db_path) as conn:
            # Delete chunks first
            conn.execute('DELETE FROM document_chunks WHERE document_id = ?', (document_id,))
            
            # Delete document
            result = conn.execute('DELETE FROM documents WHERE document_id = ?', (document_id,))
            
            return result.rowcount > 0


class RAGProcessor:
    """Main RAG processor combining document processing and retrieval"""
    
    def __init__(self, db_path: str = "rag_documents.db"):
        self.doc_processor = DocumentProcessor()
        self.rag_db = RAGDatabase(db_path)
        self.logger = logging.getLogger(f"{__name__}.RAGProcessor")
    
    async def process_and_store_document(self, file_content: bytes, filename: str, 
                                       content_type: str) -> ProcessedDocument:
        """Process and store a document for RAG"""
        self.logger.info(f"Processing document: {filename} ({content_type})")
        
        # Process document
        doc = await self.doc_processor.process_document(file_content, filename, content_type)
        
        if not doc.error_message:
            # Create chunks from the processed content
            full_content = doc.content_preview  # This is a simplified version
            chunks = self.doc_processor._create_chunks(full_content, doc.document_id)
            
            # Store in database
            self.rag_db.store_document(doc, chunks)
            
            self.logger.info(f"Stored document {doc.document_id} with {len(chunks)} chunks")
        else:
            self.logger.error(f"Failed to process document {filename}: {doc.error_message}")
        
        return doc
    
    def retrieve_context(self, query: str, max_chunks: int = 5) -> Tuple[List[DocumentChunk], str]:
        """Retrieve relevant context for a query"""
        relevant_chunks = self.rag_db.search_relevant_chunks(query, max_chunks)
        
        if not relevant_chunks:
            return [], ""
        
        # Build context string
        context_parts = []
        for chunk in relevant_chunks:
            context_parts.append(f"[Document: {chunk.document_id}]\n{chunk.content}")
        
        context = "\n\n---\n\n".join(context_parts)
        
        return relevant_chunks, context
    
    def get_document_library(self) -> List[ProcessedDocument]:
        """Get all documents in the RAG library"""
        return self.rag_db.get_all_documents()
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document from the RAG library"""
        return self.rag_db.delete_document(document_id)
    
    def get_supported_file_types(self) -> List[str]:
        """Get list of supported file types"""
        return list(self.doc_processor.supported_types.keys())